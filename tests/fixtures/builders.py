"""Programmatic test-document builders.

All PDFs and DOCX files used by tests are generated here; no third-party
documents are ever added to the repository. Fixture text stays ASCII because
PDF base-14 fonts do not cover Cyrillic.
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from dalel.ingestion.hashing import sha256_file

PAGE_TEXT = (
    "This is a synthetic environmental permit fixture page. It contains more"
    " than enough embedded text characters to count as a digital page for the"
    " OCR policy analysis."
)


def make_digital_pdf(path: Path, pages: int = 2, with_table: bool = True) -> None:
    import fitz

    doc = fitz.open()
    for index in range(pages):
        page = doc.new_page(width=595, height=842)
        page.insert_text((72, 96), f"Fixture page {index + 1}.", fontsize=14)
        page.insert_textbox(fitz.Rect(72, 120, 523, 260), PAGE_TEXT, fontsize=11)
        if with_table and index == 0:
            _draw_ruled_table(page)
    doc.save(str(path))
    doc.close()


def _draw_ruled_table(page: Any) -> None:
    import fitz

    x0, y0, x1, y1 = 72, 300, 372, 360
    rows, cols = 2, 3
    for r in range(rows + 1):
        y = y0 + (y1 - y0) * r / rows
        page.draw_line(fitz.Point(x0, y), fitz.Point(x1, y))
    for c in range(cols + 1):
        x = x0 + (x1 - x0) * c / cols
        page.draw_line(fitz.Point(x, y0), fitz.Point(x, y1))
    for r in range(rows):
        for c in range(cols):
            page.insert_text(
                (x0 + c * (x1 - x0) / cols + 6, y0 + r * (y1 - y0) / rows + 18),
                f"c{r}{c}",
                fontsize=9,
            )


def make_scanned_pdf(path: Path, pages: int = 2, render_text: str | None = None) -> None:
    """A PDF without embedded text. When ``render_text`` is given, the words are
    rasterized into a page image so OCR integration tests have something real
    to recognize."""
    import fitz

    if render_text is None:
        doc = fitz.open()
        for _ in range(pages):
            page = doc.new_page(width=595, height=842)
            page.draw_rect(fitz.Rect(100, 100, 495, 742), fill=(0.85, 0.85, 0.85))
        doc.save(str(path))
        doc.close()
        return

    source = fitz.open()
    page = source.new_page(width=595, height=842)
    page.insert_textbox(fitz.Rect(72, 96, 523, 400), render_text, fontsize=16)
    pixmap = page.get_pixmap(dpi=150)
    png_bytes = pixmap.tobytes("png")
    source.close()

    doc = fitz.open()
    for _ in range(pages):
        image_page = doc.new_page(width=595, height=842)
        image_page.insert_image(fitz.Rect(0, 0, 595, 842), stream=png_bytes)
    doc.save(str(path))
    doc.close()


def make_docx(path: Path) -> None:
    import docx

    document = docx.Document()
    document.add_heading("Fixture Summary", level=1)
    document.add_paragraph("Opening paragraph before any table.")
    document.add_heading("Emissions", level=2)
    document.add_paragraph("Emissions are described in the table below.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "substance"
    table.cell(0, 1).text = "limit"
    table.cell(1, 0).text = "dust"
    table.cell(1, 1).text = "0.5"
    document.save(str(path))


def make_fake_rar(path: Path) -> None:
    path.write_bytes(b"Rar!\x1a\x07\x00" + b"\x00" * 128)


def make_manifest_document(
    repo_root: Path,
    document_id: str,
    local_path: Path,
    document_type: str,
    role: str = "model_input",
    file_format: str | None = None,
    use_as_model_feature: bool | None = None,
    label_timing: str | None = None,
    sha256: str | None = None,
) -> dict[str, Any]:
    if use_as_model_feature is None:
        use_as_model_feature = role == "model_input"
    if label_timing is None:
        label_timing = "pre_review" if role == "model_input" else "post_review"
    return {
        "document_id": document_id,
        "local_path": str(local_path.relative_to(repo_root)),
        "original_filename": local_path.name,
        "document_type": document_type,
        "role": role,
        "use_as_model_feature": use_as_model_feature,
        "file_format": file_format or local_path.suffix.lstrip("."),
        "sha256": sha256 or sha256_file(local_path),
        "label_timing": label_timing,
        "notes": None,
    }


def write_manifest(manifest_path: Path, projects: list[dict[str, Any]]) -> None:
    manifest_path.parent.mkdir(parents=True, exist_ok=True)
    with manifest_path.open("w", encoding="utf-8") as handle:
        for project in projects:
            handle.write(json.dumps(project, ensure_ascii=False))
            handle.write("\n")
