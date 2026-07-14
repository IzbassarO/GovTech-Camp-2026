"""python-docx fallback parser for DOCX documents.

DOCX is a flow format: it has no page geometry, so page numbers and bboxes are
``null`` with explicit warnings. Provenance is paragraph/table-index based —
as precise as the format allows.
"""

from __future__ import annotations

import importlib.metadata
import logging
import re
from pathlib import Path

from dalel.ingestion.parsed import (
    OcrOutcome,
    ParsedDocument,
    ParsedImage,
    ParsedPage,
    ParsedSection,
    ParsedTable,
    SkippedTableItem,
)
from dalel.schemas.table import table_content_is_valid

logger = logging.getLogger(__name__)

PARSER_NAME = "python-docx"

_HEADING_RE = re.compile(r"^Heading\s+(\d+)$", re.IGNORECASE)


def parser_version() -> str | None:
    try:
        return importlib.metadata.version("python-docx")
    except importlib.metadata.PackageNotFoundError:
        return None


def _heading_level(style_name: str | None) -> int | None:
    if not style_name:
        return None
    match = _HEADING_RE.match(style_name.strip())
    if match:
        return int(match.group(1))
    if style_name.strip().lower() == "title":
        return 0
    return None


def parse_docx_fallback(path: Path) -> ParsedDocument:
    """Parse a DOCX with python-docx."""
    import docx  # lazy: python-docx

    document = docx.Document(str(path))
    parsed = ParsedDocument(
        parser_name=PARSER_NAME,
        parser_version=parser_version(),
        status="success",
        ocr=OcrOutcome(),
    )
    parsed.warnings.append(
        "docx flow format: page numbers and bboxes are unavailable; provenance is"
        " paragraph/table-index based"
    )

    sections: list[ParsedSection] = []
    current: ParsedSection | None = None
    preamble: list[str] = []
    all_text: list[str] = []

    def _close(section: ParsedSection | None) -> None:
        if section is not None:
            section.text = section.text.strip()
            sections.append(section)

    for paragraph in document.paragraphs:
        text = paragraph.text
        if not text.strip():
            continue
        all_text.append(text)
        style_name = paragraph.style.name if paragraph.style is not None else None
        level = _heading_level(style_name)
        if level is not None:
            _close(current)
            current = ParsedSection(
                title=text.strip(),
                level=level,
                page_start=None,
                page_end=None,
                text="",
                warnings=["docx has no page numbers"],
            )
        elif current is not None:
            current.text += text + "\n"
        else:
            preamble.append(text)
    _close(current)
    if preamble:
        sections.insert(
            0,
            ParsedSection(
                title=None,
                level=None,
                text="\n".join(preamble).strip(),
                warnings=["content before the first heading", "docx has no page numbers"],
            ),
        )
    parsed.sections = sections

    for table in document.tables:
        try:
            cells = [[cell.text for cell in row.cells] for row in table.rows]
            num_rows = len(cells)
            num_cols = max((len(row) for row in cells), default=0)
            if not table_content_is_valid(num_rows, num_cols, cells):
                parsed.skipped_empty_tables.append(
                    SkippedTableItem(
                        page_number=None,
                        reference=None,
                        extraction_method=PARSER_NAME,
                        message="docx table has no rows/columns/cell content",
                    )
                )
                continue
            parsed.tables.append(
                ParsedTable(
                    page_number=None,
                    bbox=None,
                    cells=cells,
                    num_rows=num_rows,
                    num_cols=num_cols,
                    warnings=["docx table has no page number or bbox"],
                )
            )
        except Exception as exc:
            parsed.skipped_empty_tables.append(
                SkippedTableItem(
                    page_number=None,
                    reference=None,
                    extraction_method=PARSER_NAME,
                    message=f"docx table extraction failed: {exc}",
                )
            )

    try:
        for rel in document.part.rels.values():
            if "image" not in rel.reltype:
                continue
            try:
                blob = rel.target_part.blob
            except Exception as exc:
                parsed.images.append(ParsedImage(warnings=[f"docx image blob unavailable: {exc}"]))
                continue
            parsed.images.append(
                ParsedImage(
                    page_number=None,
                    bbox=None,
                    png_bytes=blob,
                    warnings=["docx image has no page number or bbox"],
                )
            )
    except Exception as exc:
        parsed.warnings.append(f"docx image enumeration failed: {exc}")

    # Flow document: represent the full text as a single pseudo-page so that
    # downstream consumers always find text in pages.jsonl. Page geometry is null.
    full_text = "\n".join(all_text)
    parsed.pages.append(
        ParsedPage(
            page_number=1,
            width=None,
            height=None,
            rotation=None,
            text=full_text,
            ocr_applied=False,
            has_embedded_text=bool(full_text.strip()),
            warnings=["pseudo-page: docx flow text without page geometry"],
        )
    )
    return parsed
